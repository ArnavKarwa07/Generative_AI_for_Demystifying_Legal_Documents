import os
import json
import uuid
import base64
import asyncio
from datetime import datetime
from typing import Dict, List, Optional, Any
import aiofiles
import shutil
from docx import Document as DocxDocument
from sqlalchemy.orm import Session
from app.database import get_db
from app.models import FileStorage as FileStorageModel


class FileStorage:
    def __init__(self):
        self.documents_dir = "documents"
        self.uploads_dir = "uploads"
        self.drafts_dir = "drafts"
        self.metadata_file = os.path.join(self.documents_dir, "metadata.json")

        # Create directories for backward compatibility
        os.makedirs(self.documents_dir, exist_ok=True)
        os.makedirs(self.uploads_dir, exist_ok=True)
        os.makedirs(self.drafts_dir, exist_ok=True)

        # Initialize metadata file if it doesn't exist
        if not os.path.exists(self.metadata_file):
            self._save_metadata({})

    def _load_metadata(self) -> Dict[str, Any]:
        """Load metadata from JSON file"""
        try:
            with open(self.metadata_file, "r") as f:
                return json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            return {}

    def _save_metadata(self, metadata: Dict[str, Any]):
        """Save metadata to JSON file"""
        with open(self.metadata_file, "w") as f:
            json.dump(metadata, f, indent=2, default=str)

    def _get_db_session(self) -> Session:
        """Get database session"""
        return next(get_db())

    async def save_file_to_db(
        self,
        file_content: bytes,
        filename: str,
        content_type: str,
        user_id: Optional[int] = None,
    ) -> str:
        """Save file to database using base64 encoding"""
        file_id = str(uuid.uuid4())

        # Encode file content to base64
        encoded_content = base64.b64encode(file_content).decode("utf-8")

        # Save to database
        db = self._get_db_session()
        try:
            file_record = FileStorageModel(
                id=file_id,
                filename=filename,
                content_type=content_type,
                file_data=encoded_content,
                file_size=len(file_content),
                created_by_id=user_id,
            )
            db.add(file_record)
            db.commit()
            db.refresh(file_record)
            return file_id
        except Exception as e:
            db.rollback()
            raise e
        finally:
            db.close()

    async def get_file_from_db(self, file_id: str) -> Optional[Dict[str, Any]]:
        """Get file from database"""
        db = self._get_db_session()
        try:
            file_record = (
                db.query(FileStorageModel)
                .filter(FileStorageModel.id == file_id)
                .first()
            )
            if not file_record:
                return None

            # Decode base64 content
            file_content = base64.b64decode(file_record.file_data.encode("utf-8"))

            return {
                "id": file_record.id,
                "filename": file_record.filename,
                "content_type": file_record.content_type,
                "file_content": file_content,
                "file_size": file_record.file_size,
                "created_at": file_record.created_at,
                "created_by_id": file_record.created_by_id,
            }
        finally:
            db.close()

    async def save_document(
        self,
        file_content: bytes,
        filename: str,
        file_type: Optional[str] = None,
        analysis: Optional[Dict[str, Any]] = None,
        uploaded_at: Optional[str] = None,
        user_id: Optional[int] = None,
    ) -> str:
        """Save uploaded document to database and return document ID"""
        content_type = file_type or f"application/{filename.split('.')[-1].lower()}"
        doc_id = await self.save_file_to_db(
            file_content, filename, content_type, user_id
        )

        # Also save legacy metadata for backward compatibility
        metadata = self._load_metadata()
        metadata[doc_id] = {
            "id": doc_id,
            "filename": filename,
            "file_type": file_type or filename.split(".")[-1].lower(),
            "uploaded_at": uploaded_at or datetime.now().isoformat(),
            "storage_type": "database",
            "analysis": analysis or {},
        }
        self._save_metadata(metadata)

        return doc_id

    def get_document(self, doc_id: str) -> Optional[Dict[str, Any]]:
        """Get document metadata by ID"""
        # Try to get from database first
        try:
            file_data = asyncio.run(self.get_file_from_db(doc_id))
            if file_data:
                return {
                    "id": file_data["id"],
                    "filename": file_data["filename"],
                    "file_type": file_data["content_type"],
                    "uploaded_at": (
                        file_data["created_at"].isoformat()
                        if file_data["created_at"]
                        else None
                    ),
                    "storage_type": "database",
                    "file_size": file_data["file_size"],
                }
        except:
            pass

        # Fallback to metadata file
        metadata = self._load_metadata()
        return metadata.get(doc_id)

    def get_all_documents(self) -> List[Dict[str, Any]]:
        """Get all documents metadata"""
        documents = []

        # Get from database
        db = self._get_db_session()
        try:
            db_files = db.query(FileStorageModel).all()
            for file_record in db_files:
                documents.append(
                    {
                        "id": file_record.id,
                        "filename": file_record.filename,
                        "file_type": file_record.content_type,
                        "uploaded_at": (
                            file_record.created_at.isoformat()
                            if file_record.created_at is not None
                            else None
                        ),
                        "storage_type": "database",
                        "file_size": file_record.file_size,
                    }
                )
        finally:
            db.close()

        # Also include legacy metadata
        metadata = self._load_metadata()
        for doc_id, doc_data in metadata.items():
            if doc_data.get("storage_type") != "database":  # Avoid duplicates
                documents.append(doc_data)

        return documents

    def delete_document(self, doc_id: str) -> bool:
        """Delete document from database and metadata"""
        success = False

        # Try to delete from database
        db = self._get_db_session()
        try:
            file_record = (
                db.query(FileStorageModel).filter(FileStorageModel.id == doc_id).first()
            )
            if file_record:
                db.delete(file_record)
                db.commit()
                success = True
        except:
            db.rollback()
        finally:
            db.close()

        # Also remove from legacy metadata
        metadata = self._load_metadata()
        if doc_id in metadata:
            # Remove directory if it exists
            doc_dir = os.path.join(self.documents_dir, doc_id)
            if os.path.exists(doc_dir):
                shutil.rmtree(doc_dir)

            del metadata[doc_id]
            self._save_metadata(metadata)
            success = True

        return success

    async def update_document_analysis(self, doc_id: str, analysis: Dict[str, Any]):
        """Update document analysis"""
        metadata = self._load_metadata()
        if doc_id in metadata:
            metadata[doc_id]["analysis"] = analysis
            metadata[doc_id]["analyzed_at"] = datetime.now().isoformat()
            self._save_metadata(metadata)

    async def save_draft(self, content: str, title: Optional[str] = None) -> str:
        """Save draft document"""
        draft_id = str(uuid.uuid4())
        draft_path = os.path.join(self.drafts_dir, f"{draft_id}.md")

        async with aiofiles.open(draft_path, "w") as f:
            await f.write(content)

        # Save draft metadata
        metadata = self._load_metadata()
        if "drafts" not in metadata:
            metadata["drafts"] = {}

        metadata["drafts"][draft_id] = {
            "id": draft_id,
            "title": title or f"Draft {draft_id[:8]}",
            "created_at": datetime.now().isoformat(),
            "file_path": draft_path,
        }
        self._save_metadata(metadata)

        return draft_id

    async def save_draft_docx(
        self, content: str, title: Optional[str] = None, user_id: Optional[int] = None
    ) -> Dict[str, Any]:
        """Create and save a draft as a DOCX file and record metadata.

        Returns a dict with keys: id, filename, file_path, title, created_at
        """
        draft_id = str(uuid.uuid4())
        filename = f"{draft_id}.docx"

        # Build DOCX file in memory
        doc = DocxDocument()
        if title:
            doc.add_heading(title, level=1)
        # Split content by double newlines into paragraphs; preserve basic line breaks
        if content:
            for block in str(content).split("\n\n"):
                # Replace single newlines within a block with line breaks
                p = doc.add_paragraph()
                lines = block.split("\n")
                for i, line in enumerate(lines):
                    if i == 0:
                        p.add_run(line)
                    else:
                        p.add_run()  # break by starting a new run
                        p.add_run("\n" + line)
        else:
            doc.add_paragraph("")

        # Save to memory buffer
        import io

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_content = buffer.getvalue()

        # Save to database
        file_id = await self.save_file_to_db(
            docx_content,
            filename,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            user_id,
        )

        created_at = datetime.now().isoformat()

        # Save draft metadata for backward compatibility
        metadata = self._load_metadata()
        if "drafts" not in metadata:
            metadata["drafts"] = {}

        metadata["drafts"][file_id] = {
            "id": file_id,
            "title": title or f"Draft {file_id[:8]}",
            "created_at": created_at,
            "filename": filename,
            "file_type": "docx",
            "storage_type": "database",
        }
        self._save_metadata(metadata)

        return {
            "id": file_id,
            "title": title or f"Draft {file_id[:8]}",
            "created_at": created_at,
            "filename": filename,
        }

    def get_draft(self, draft_id: str) -> Optional[str]:
        """Get draft content"""
        metadata = self._load_metadata()
        drafts = metadata.get("drafts", {})

        if draft_id not in drafts:
            return None

        draft_path = drafts[draft_id]["file_path"]
        try:
            with open(draft_path, "r") as f:
                return f.read()
        except FileNotFoundError:
            return None

    def get_all_drafts(self) -> List[Dict[str, Any]]:
        """Get all drafts metadata"""
        metadata = self._load_metadata()
        return list(metadata.get("drafts", {}).values())


# Global storage instance
file_storage = FileStorage()
